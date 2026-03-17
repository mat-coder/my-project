"""
Financial Status Certificate — Template Filler
Run with:  streamlit run app.py
"""

import streamlit as st
from docx import Document
from io import BytesIO
from copy import deepcopy


TEMPLATE_PATH = "Mohd_Aslam__1_.docx"

# ── Core helpers ──────────────────────────────────────────────────────────────

def replace_in_runs(para, old: str, new: str):
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""


def replace_all(doc: Document, old: str, new: str):
    for para in doc.paragraphs:
        replace_in_runs(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, old, new)


def fill_family_table(table, rows):
    current_data_rows = len(table.rows) - 1
    
    # Add rows if needed
    while current_data_rows < len(rows):
        last_row = table.rows[-1]
        new_tr = deepcopy(last_row._tr)
        last_row._tr.addnext(new_tr)
        current_data_rows += 1
        
    # Remove rows if needed
    while current_data_rows > len(rows):
        row = table.rows[-1]
        row._element.getparent().remove(row._element)
        current_data_rows -= 1

    for i, rd in enumerate(rows):
        ri = i + 1
        cells = table.rows[ri].cells
        
        # Update S.No
        if len(cells) > 0:
            para = cells[0].paragraphs[0]
            val = str(i + 1)
            if para.runs:
                para.runs[0].text = val
                for r in para.runs[1:]:
                    r.text = ""
            else:
                para.add_run(val)

        pairs = [(1, rd["name"]), (2, rd["age"]), (3, rd["relation"]),
                 (4, rd["marital"]), (5, rd["occupation"]), (6, rd["education"]), (7, rd["income"])]
        for col_idx, value in pairs:
            if col_idx < len(cells):
                para = cells[col_idx].paragraphs[0]
                if para.runs:
                    para.runs[0].text = value
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(value)


def fill_enquiry_table(table, pension, movable, other_income, fin_pos, remarks):
    values = [pension, movable, other_income, fin_pos, remarks]
    for ri, val in enumerate(values):
        if ri < len(table.rows):
            cells = table.rows[ri].cells
            if len(cells) >= 2:
                para = cells[1].paragraphs[0]
                if para.runs:
                    para.runs[0].text = val
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(val)


def validate_inputs(inputs: dict, family_rows: list) -> list:
    errors = []
    
    # Validation for top level inputs being non-empty
    required_keys = {
        'lr_no': 'Lr. No. (Certificate)', 'note_lr_no': 'Lr. No. (Note File)', 'dated': 'Dated',
        'tahsildar_lr': 'Tahsildar Lr. No.', 'tahsildar_dt': 'Tahsildar Lr. Dated', 'mandal': 'Mandal',
        'applicant_name': 'Applicant Full Name', 'deceased_name': 'Deceased Name', 'address': 'Full Address',
        'dec_father': "Deceased's Father Name", 'designation': 'Designation / Post',
        'old_office': 'Old Office Name', 'new_office': 'Renamed / Current Office Name',
        'date_of_death': 'Date of Death', 'pension': 'Pension Receiving', 'movable': 'Movable Properties',
        'other_income': 'Other Source of Income', 'remarks': 'Remarks'
    }
    
    for key, label in required_keys.items():
        if not inputs.get(key, "").strip():
            errors.append(f"'{label}' cannot be empty.")
            
    # Validation for family rows
    for i, row in enumerate(family_rows):
        if not row['name'].strip():
            errors.append(f"Family Member {i+1}: Name cannot be empty.")
        if not row['age'].strip():
            errors.append(f"Family Member {i+1}: Age cannot be empty.")
        elif not row['age'].strip().isdigit():
            errors.append(f"Family Member {i+1}: Age must be a valid positive number.")
        if not row['relation'].strip():
            errors.append(f"Family Member {i+1}: Relation cannot be empty.")
            
    return errors


def generate_doc(inputs: dict, family_rows: list) -> bytes:
    doc = Document(TEMPLATE_PATH)
    f = inputs

    replacements = [
        ("Lr.No.G/478/2026",   f"Lr.No.{f['lr_no']}"),
        ("Lr.No.G/478/2025",   f"Lr.No.{f['note_lr_no']}"),
        ("Dated.    .03.2026.", f"Dated.    {f['dated']}."),
        ("Lr. No. C/390/2026",  f"Lr. No. {f['tahsildar_lr']}"),
        ("Dated.28.02.2026",    f"Dated.{f['tahsildar_dt']}"),
        ("Balapur Mandal",      f['mandal']),
        ("Sri. Mohd Aslam  S/o",      f"Sri. {f['applicant_name']} {f['applicant_rel']}"),
        ("Sri. Mohammad Aslam  S/o",  f"Sri. {f['applicant_name']} {f['applicant_rel']}"),
        ("Mohd Aslam  S/o",           f"{f['applicant_name']} {f['applicant_rel']}"),
        ("Mohammad Aslam  S/o",       f"{f['applicant_name']} {f['applicant_rel']}"),
        ("Sri. Mohammad Aslam",       f"Sri. {f['applicant_name']}"),
        ("Sri. Mohd Aslam",           f"Sri. {f['applicant_name']}"),
        ("Mohammad Aslam",            f['applicant_name']),
        ("Late. Mohammad Akbar",      f"Late. {f['deceased_name']}"),
        ("Late Mohammad Akbar",       f"Late {f['deceased_name']}"),
        ("Late Mohammad Haji",        f"Late {f['dec_father']}"),
        ("H. No. 9-221-89/812, Metro City-IV, Balapur Village and Mandal, Ranga Reddy District",
         f['address']),
        ("Record Assistant",          f['designation']),
        ("Erstwhile O/o  The Executive Engineer, Kaleshwaram Project, Construction Division No.7, Gajwel",
         f['old_office']),
        ("Irrigation Division No.6, Gajwel, Siddipet District, Telangana State",
         f['new_office']),
        ("12.07.2025",                f['date_of_death']),
    ]

    for old, new in replacements:
        replace_all(doc, old, new)

    family_tables, enquiry_tables = [], []
    for table in doc.tables:
        header_text = " ".join(c.text for c in table.rows[0].cells) if table.rows else ""
        if "Name of the Individual" in header_text:
            family_tables.append(table)
        else:
            enquiry_tables.append(table)

    for ft in family_tables:
        fill_family_table(ft, family_rows)
    for et in enquiry_tables:
        fill_enquiry_table(et, f['pension'], f['movable'],
                           f['other_income'], f['fin_pos'], f['remarks'])

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Streamlit UI ──────────────────────────────────────────────────────────────

st.set_page_config(page_title="Financial Status Certificate Generator",
                   page_icon="📄", layout="wide")

st.title("📄 Financial Status Certificate — Template Filler")
st.caption("Fill in the fields below. Original document layout and formatting are fully preserved.")

st.header("1️⃣  Letter & Reference Details")
c1, c2, c3 = st.columns(3)
lr_no        = c1.text_input("Lr. No. (Certificate)", value="G/478/2026")
note_lr_no   = c2.text_input("Lr. No. (Note File)",   value="G/478/2025")
dated        = c3.text_input("Dated",                  value=".03.2026")
c4, c5, c6 = st.columns(3)
tahsildar_lr = c4.text_input("Tahsildar Lr. No.",      value="C/390/2026")
tahsildar_dt = c5.text_input("Tahsildar Lr. Dated",    value="28.02.2026")
mandal       = c6.text_input("Mandal",                 value="Balapur Mandal")

st.header("2️⃣  Applicant Details")
c7, c8, c9 = st.columns(3)
applicant_name = c7.text_input("Applicant Full Name",    value="Mohd Aslam")
applicant_rel  = c8.selectbox("Relation to Deceased",    ["S/o", "D/o", "W/o"], index=0)
deceased_name  = c9.text_input("Deceased Name",          value="Mohammad Akbar")
address = st.text_input("Full Address",
    value="H. No. 9-221-89/812, Metro City-IV, Balapur Village and Mandal, Ranga Reddy District")

st.header("3️⃣  Deceased Employee Details")
c10, c11 = st.columns(2)
dec_father    = c10.text_input("Deceased's Father Name", value="Mohammad Haji")
designation   = c11.text_input("Designation / Post",    value="Record Assistant")
old_office = st.text_input("Old Office Name (Erstwhile)",
    value="O/o The Executive Engineer, Kaleshwaram Project, Construction Division No.7, Gajwel")
new_office = st.text_input("Renamed / Current Office Name",
    value="Irrigation Division No.6, Gajwel, Siddipet District, Telangana State")
date_of_death = st.text_input("Date of Death", value="12.07.2025")

st.header("4️⃣  Further Enquiry Details")
c12, c13 = st.columns(2)
pension      = c12.text_input("Pension Receiving?",          value="Not Receiving")
movable      = c13.text_input("Movable Properties / House",  value="Own House")
c14, c15, c16 = st.columns(3)
other_income = c14.text_input("Any Other Source of Income",  value="Nil")
fin_pos      = c15.selectbox("Overall Financial Position",   ["Sound", "Unsound", "Average"], index=1)
remarks      = c16.text_input("Any Other Remarks",           value="Nil")

st.header("5️⃣  Family Members")

DEFAULT_ROWS = [
    {"name": "Shaheda Begum",   "age": "55", "relation": "Wife",     "marital": "Widow",
     "occupation": "House wife",       "education": "Nil",          "income": "Nil"},
    {"name": "Mohd Aslam",      "age": "34", "relation": "Son",      "marital": "Married",
     "occupation": "Un-Employee",      "education": "Intermediate", "income": "Nil"},
    {"name": "Mohd Akram",      "age": "33", "relation": "Son",      "marital": "Married",
     "occupation": "Private Employee", "education": "Intermediate", "income": "18,000/-"},
    {"name": "Ayesha Tabassum", "age": "30", "relation": "Daughter", "marital": "Married",
     "occupation": "House wife",       "education": "Intermediate", "income": "--"},
]

if "num_members" not in st.session_state:
    st.session_state.num_members = 4

def add_member():
    st.session_state.num_members += 1

def remove_member():
    if st.session_state.num_members > 1:
        st.session_state.num_members -= 1

family_rows = []
for i in range(st.session_state.num_members):
    d = DEFAULT_ROWS[i] if i < len(DEFAULT_ROWS) else {}
    with st.expander(f"Member {i + 1}", expanded=True):
        mc1, mc2, mc3, mc4 = st.columns(4)
        name     = mc1.text_input("Name",          value=d.get("name", ""),       key=f"name_{i}")
        age      = mc2.text_input("Age",            value=d.get("age", ""),        key=f"age_{i}")
        relation = mc3.text_input("Relation",       value=d.get("relation", ""),   key=f"rel_{i}")
        marital  = mc4.text_input("Marital Status", value=d.get("marital", ""),    key=f"mar_{i}")
        mc5, mc6, mc7 = st.columns(3)
        occ = mc5.text_input("Occupation",     value=d.get("occupation", ""), key=f"occ_{i}")
        edu = mc6.text_input("Education",      value=d.get("education", ""),  key=f"edu_{i}")
        inc = mc7.text_input("Income / Month", value=d.get("income", ""),     key=f"inc_{i}")
        family_rows.append({"name": name, "age": age, "relation": relation, "marital": marital,
                             "occupation": occ, "education": edu, "income": inc})

col_add, col_rem, _ = st.columns([1, 1, 4])
col_add.button("➕ Add Family Member", on_click=add_member)
col_rem.button("➖ Remove Member", on_click=remove_member)

st.divider()

if st.button("🖨️  Generate Filled Certificate (.docx)", type="primary", use_container_width=True):
    inputs = dict(
        lr_no=lr_no, note_lr_no=note_lr_no, dated=dated,
        tahsildar_lr=tahsildar_lr, tahsildar_dt=tahsildar_dt, mandal=mandal,
        applicant_name=applicant_name, applicant_rel=applicant_rel,
        deceased_name=deceased_name, address=address,
        dec_father=dec_father, designation=designation,
        old_office=old_office, new_office=new_office, date_of_death=date_of_death,
        pension=pension, movable=movable, other_income=other_income,
        fin_pos=fin_pos, remarks=remarks,
    )
    try:
        errors = validate_inputs(inputs, family_rows)
        if errors:
            for error in errors:
                st.error(f"❌ {error}")
        else:
            doc_bytes = generate_doc(inputs, family_rows)
            st.success("✅ Document generated successfully!")
            st.download_button(
                label="⬇️  Download Filled Certificate",
                data=doc_bytes,
                file_name="Financial_Status_Certificate_Filled.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
    except Exception as e:
        st.error(f"❌ Error: {e}")
        st.exception(e)

st.caption("Government of Telangana — Revenue Department — Financial Status Certificate")
