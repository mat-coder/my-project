"""
NOTE FILE — Birth/Death/Caste Certificate Template Filler
Run with:  streamlit run app2.py
"""

import streamlit as st
from docx import Document
from io import BytesIO


TEMPLATE_PATH = "docfile2.docx"


# ── Core helpers ──────────────────────────────────────────────────────────────

def replace_in_runs(para, old: str, new: str):
    """Replace `old` with `new` across the joined run text of a paragraph,
    preserving the first run's formatting."""
    full = "".join(r.text for r in para.runs)
    if old not in full:
        return
    full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = full
        for r in para.runs[1:]:
            r.text = ""


def replace_all(doc: Document, old: str, new: str):
    """Apply a find-replace across every paragraph and every table cell."""
    for para in doc.paragraphs:
        replace_in_runs(para, old, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_runs(para, old, new)


def generate_doc(inputs: dict) -> bytes:
    """Open the template, apply all placeholder replacements, return bytes."""
    doc = Document(TEMPLATE_PATH)
    f = inputs

    # Determine the selected certificate type label
    cert_type = f["cert_type"]  # e.g. "Birth", "Death", or "Caste"

    # Build replacement pairs  (template text → filled text)
    replacements = [
        # --- Header ---
        ("Lr.No.G/______",              f"Lr.No.G/{f['lr_no']}"),
        ("Date:     -     -2026",       f"Date: {f['date_day']}-{f['date_month']}-{f['date_year']}"),

        # --- Subject table (Row 0, Cell 1) ---
        ("Birth/Death/Caste Certificate",  f"{cert_type} Certificate"),
        ("A/o ___________________________", f"A/o {f['applicant_name']}"),

        # --- Reference table (Row 1) ---
        ("Tahsildar _______________",    f"Tahsildar {f['tahsildar_mandal']}"),
        ("Lr. No.___________",          f"Lr. No. {f['ref_lr_no']}"),
        ("Dated______________",         f"Dated {f['ref_dated']}"),

        # --- Body paragraph (P6) ---
        ("Tahsildar __________________", f"Tahsildar {f['tahsildar_mandal']}"),
        ("Sri._________________________", f"Sri. {f['applicant_name']}"),
        ("W/o,S/o,D/o__________________________",
         f"{f['relation']} {f['father_husband_name']}"),
        (",R/o__________________________________",
         f", R/o {f['address']}"),
        ("________________Mandal",       f"{f['mandal']} Mandal"),
        ("LRBD/CND No.___________________________________________.",
         f"LRBD/CND No. {f['lrbd_cnd_no']}."),
        ("The applicant caste is._____________________________",
         f"The applicant caste is. {f['caste']}"),

        # --- P7: certificate type in approval line ---
        ("Birth/Death/Caste certificate", f"{cert_type} certificate"),

        # --- P8: Cause of rejection ---
        ("Cause of Rejection.\t",
         f"Cause of Rejection. {f['cause_of_rejection']}\t"
         if f['cause_of_rejection'].strip() else "Cause of Rejection.\t"),
    ]

    for old, new in replacements:
        replace_all(doc, old, new)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Streamlit UI ──────────────────────────────────────────────────────────────

st.set_page_config(
    page_title="NOTE FILE — Certificate Template Filler",
    page_icon="📋",
    layout="wide",
)

# ── Custom CSS for a polished, premium look ──────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');

html, body, [class*="st-"] {
    font-family: 'Inter', sans-serif;
}

/* Main container */
.block-container {
    max-width: 900px;
    padding-top: 2rem;
}

/* Header gradient */
.main-header {
    background: linear-gradient(135deg, #1e3a5f 0%, #2d6a9f 50%, #4ea8de 100%);
    color: white;
    padding: 1.8rem 2rem;
    border-radius: 14px;
    margin-bottom: 1.5rem;
    box-shadow: 0 4px 20px rgba(30, 58, 95, 0.25);
}
.main-header h1 {
    margin: 0; font-size: 1.6rem; font-weight: 700;
}
.main-header p {
    margin: 0.3rem 0 0; opacity: 0.85; font-size: 0.95rem;
}

/* Section cards */
.section-card {
    background: #f8fafc;
    border: 1px solid #e2e8f0;
    border-radius: 12px;
    padding: 1.4rem 1.6rem 1rem;
    margin-bottom: 1.2rem;
    transition: box-shadow 0.2s;
}
.section-card:hover {
    box-shadow: 0 2px 12px rgba(0,0,0,0.06);
}
.section-card h3 {
    margin: 0 0 0.8rem;
    font-size: 1.05rem;
    color: #1e3a5f;
    font-weight: 600;
}

/* Footer */
.footer-text {
    text-align: center;
    color: #94a3b8;
    font-size: 0.82rem;
    margin-top: 2rem;
    padding-bottom: 1rem;
}
</style>
""", unsafe_allow_html=True)

# ── Header ───────────────────────────────────────────────────────────────────
st.markdown("""
<div class="main-header">
    <h1>📋 NOTE FILE — Certificate Template Filler</h1>
    <p>Fill in the fields below. The original document layout and formatting are fully preserved.</p>
</div>
""", unsafe_allow_html=True)

# ── 1. Letter & Reference Details ────────────────────────────────────────────
st.markdown('<div class="section-card"><h3>1️⃣ Letter & Reference Details</h3></div>', unsafe_allow_html=True)

c1, c2 = st.columns(2)
lr_no       = c1.text_input("Lr. No. (number after G/)", placeholder="e.g. 478")
cert_type   = c2.selectbox("Certificate Type", ["Birth", "Death", "Caste"])

c3, c4, c5 = st.columns(3)
date_day    = c3.text_input("Date — Day",   placeholder="e.g. 15")
date_month  = c4.text_input("Date — Month", placeholder="e.g. 03")
date_year   = c5.text_input("Date — Year",  value="2026")

# ── 2. Tahsildar Reference ───────────────────────────────────────────────────
st.markdown('<div class="section-card"><h3>2️⃣ Tahsildar Reference</h3></div>', unsafe_allow_html=True)

c6, c7, c8 = st.columns(3)
tahsildar_mandal = c6.text_input("Tahsildar Mandal", placeholder="e.g. Balapur")
ref_lr_no        = c7.text_input("Reference Lr. No.", placeholder="e.g. C/390/2026")
ref_dated         = c8.text_input("Reference Dated",  placeholder="e.g. 28.02.2026")

# ── 3. Applicant Details ─────────────────────────────────────────────────────
st.markdown('<div class="section-card"><h3>3️⃣ Applicant Details</h3></div>', unsafe_allow_html=True)

c9, c10, c11 = st.columns(3)
applicant_name      = c9.text_input("Applicant Full Name", placeholder="e.g. Mohd Aslam")
relation            = c10.selectbox("Relation", ["S/o", "D/o", "W/o"])
father_husband_name = c11.text_input("Father / Husband Name", placeholder="e.g. Mohammad Akbar")

c12, c13 = st.columns(2)
address = c12.text_input("Address / Village", placeholder="e.g. Balapur Village")
mandal  = c13.text_input("Mandal", placeholder="e.g. Balapur")

# ── 4. Additional Details ────────────────────────────────────────────────────
st.markdown('<div class="section-card"><h3>4️⃣ Additional Details</h3></div>', unsafe_allow_html=True)

c14, c15 = st.columns(2)
lrbd_cnd_no = c14.text_input("LRBD / CND No.", placeholder="e.g. 12345/2026")
caste       = c15.text_input("Applicant Caste", placeholder="e.g. OC")

cause_of_rejection = st.text_area(
    "Cause of Rejection (leave blank if none)",
    height=80,
    placeholder="Enter cause of rejection, if any…",
)

# ── Generate Button ──────────────────────────────────────────────────────────
st.divider()

if st.button("🖨️  Generate Filled NOTE FILE (.docx)", type="primary", use_container_width=True):
    inputs = dict(
        lr_no=lr_no, cert_type=cert_type,
        date_day=date_day, date_month=date_month, date_year=date_year,
        tahsildar_mandal=tahsildar_mandal, ref_lr_no=ref_lr_no, ref_dated=ref_dated,
        applicant_name=applicant_name, relation=relation,
        father_husband_name=father_husband_name,
        address=address, mandal=mandal,
        lrbd_cnd_no=lrbd_cnd_no, caste=caste,
        cause_of_rejection=cause_of_rejection,
    )

    # Quick validation — make sure the critical fields are filled
    required = {
        "lr_no": "Lr. No.",
        "date_day": "Date — Day",
        "date_month": "Date — Month",
        "applicant_name": "Applicant Full Name",
        "father_husband_name": "Father / Husband Name",
        "address": "Address / Village",
        "mandal": "Mandal",
        "tahsildar_mandal": "Tahsildar Mandal",
        "ref_lr_no": "Reference Lr. No.",
        "ref_dated": "Reference Dated",
        "lrbd_cnd_no": "LRBD / CND No.",
        "caste": "Applicant Caste",
    }

    errors = [f"'{label}' cannot be empty." for key, label in required.items()
              if not inputs.get(key, "").strip()]

    if errors:
        for e in errors:
            st.error(f"❌ {e}")
    else:
        try:
            doc_bytes = generate_doc(inputs)
            st.success("✅ Document generated successfully!")
            st.download_button(
                label="⬇️  Download Filled NOTE FILE",
                data=doc_bytes,
                file_name="NOTE_FILE_Filled.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as exc:
            st.error(f"❌ Error generating document: {exc}")
            st.exception(exc)

# ── Footer ───────────────────────────────────────────────────────────────────
st.markdown(
    '<div class="footer-text">Government of Telangana — Revenue Department — Certificate Note File</div>',
    unsafe_allow_html=True,
)
