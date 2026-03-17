import unittest
from io import BytesIO
from docx import Document
import app
from app import generate_doc, TEMPLATE_PATH
import os

class TestAppDocGeneration(unittest.TestCase):
    def setUp(self):
        # We need to make sure the template exists for tests to run
        self.assertTrue(os.path.exists(TEMPLATE_PATH), f"Template {TEMPLATE_PATH} missing.")

        self.sample_inputs = dict(
            lr_no="G/999/2026", note_lr_no="G/888/2025", dated="15.03.2026",
            tahsildar_lr="C/777/2026", tahsildar_dt="10.02.2026", mandal="Test Mandal",
            applicant_name="Test Applicant", applicant_rel="S/o",
            deceased_name="Test Deceased", address="Test Address Line",
            dec_father="Test Father", designation="Senior Tester",
            old_office="Test Old Office", new_office="Test New Office", date_of_death="01.01.2025",
            pension="Receiving", movable="No House", other_income="1000",
            fin_pos="Average", remarks="None",
        )
        self.sample_family_rows = [
            {"name": "Wife Name", "age": "50", "relation": "Wife", "marital": "Widow", "occupation": "Home", "education": "Degree", "income": "0"},
            {"name": "Son 1", "age": "25", "relation": "Son", "marital": "Single", "occupation": "Student", "education": "B.Tech", "income": "0"},
        ]

    def test_generate_doc_executes_without_error(self):
        """Test that document generation doesn't throw exceptions with standard data."""
        try:
            doc_bytes = generate_doc(self.sample_inputs, self.sample_family_rows)
            self.assertIsInstance(doc_bytes, bytes)
            self.assertGreater(len(doc_bytes), 0)
        except Exception as e:
            self.fail(f"generate_doc raised an exception: {e}")

    def test_dynamic_family_rows_expansion(self):
        """Test adding more rows than the default template structure handles dynamically."""
        extended_family_rows = self.sample_family_rows + [
            {"name": f"Child {i}", "age": str(20-i), "relation": "Child", "marital": "Single", "occupation": "Student", "education": "School", "income": "0"}
            for i in range(10)
        ]
        
        doc_bytes = generate_doc(self.sample_inputs, extended_family_rows)
        
        # Load the generated document to verify it contains the generated rows.
        doc = Document(BytesIO(doc_bytes))
        family_tables = [t for t in doc.tables if "Name of the Individual" in (" ".join(c.text for c in t.rows[0].cells) if t.rows else "")]
        
        self.assertTrue(len(family_tables) > 0, "Could not find family tables in output doc")
        
        for table in family_tables:
            # Table rows = Header + data rows + potentially footer. In this format, each member should be in the table.
            # We verify the last row has "Child 9"
            found_child_9 = False
            for row in table.rows:
                row_text = "".join(cell.text for cell in row.cells)
                if "Child 9" in row_text:
                    found_child_9 = True
                    break
            self.assertTrue(found_child_9, "Dynamic rows were not properly injected into the table structure.")

    def test_dynamic_family_rows_reduction(self):
        """Test reducing the number of rows than what is implicitly in the template."""
        reduced_family_rows = [
            {"name": "Only Child", "age": "20", "relation": "Son", "marital": "Single", "occupation": "Student", "education": "College", "income": "0"}
        ]
        
        doc_bytes = generate_doc(self.sample_inputs, reduced_family_rows)
        doc = Document(BytesIO(doc_bytes))
        family_tables = [t for t in doc.tables if "Name of the Individual" in (" ".join(c.text for c in t.rows[0].cells) if t.rows else "")]
        
        for table in family_tables:
            for row in table.rows[1:]:
                row_text = "".join(cell.text for cell in row.cells)
                self.assertNotIn("Mohd Akram", row_text, "A default name from template was not cleared/removed properly.")


if __name__ == '__main__':
    unittest.main()
